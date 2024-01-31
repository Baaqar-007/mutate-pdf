# from pdf2docx import Converter
# from docx2pdf import convert
# from docx import Document

# def bold_first_two_letters(word):
#     if len(word) >= 2:
#         return "**" + word[:2] + "**" + word[2:]
#     else:
#         return word


# # Convert PDF to Word
# cv = Converter('658a1edd4ad6601ebd47bd3e_FeelGoodProductivity_Introduction.pdf')
# cv.convert('output.docx', start=0, end=None)
# cv.close()

# # Open the converted Word file
# doc = Document('output.docx')

# # Modify the text
# # for paragraph in doc.paragraphs:
# #     words = paragraph.text.split()
# #     modified_words = [bold_first_two_letters(word) for word in words]
# #     paragraph.text = ' '.join(modified_words)
# for para in doc.paragraphs:
#         # Clear the paragraph text
#         words = para.text.split()
#         for word in words:
#             # Add the first two letters as bold
#             if len(word) >= 2:
#                 para.add_run(word[:2]).bold = True
#                 para.add_run(word[2:] + " ")
#             else:
#                 para.add_run(word + " ")

# # Save the modified Word file
# doc.save('output.docx')
# # Convert the modified Word file back to PDF
# convert("output.docx","output.pdf")
from pdf2docx import Converter
from docx2pdf import convert
from docx import Document

def bold_first_two_letters(word):
    if len(word) >= 2:
        return word[:2], word[2:]
    else:
        return word, ''

# Convert PDF to Word
cv = Converter('input_file.pdf')
cv.convert('output.docx', start=0, end=None)
cv.close()

# Open the converted Word file
doc = Document('output.docx')

# Modify the text
for para in doc.paragraphs:
    # Clear the paragraph text
    words = para.text.split()
    para.clear()

    for word in words:
        # Add the first two letters as bold
        bold_part, rest_part = bold_first_two_letters(word)
        para.add_run(bold_part).bold = True
        para.add_run(rest_part + " ")

# Save the modified Word file
doc.save('output.docx')

# Convert the modified Word file back to PDF
convert("output.docx","output.pdf")

